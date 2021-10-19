import pandas as pd
import csv
import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Mm
from docx.oxml.shared import OxmlElement
from tqdm import tqdm
import os


filename="chat.txt"

def _convert_to_csv(filename):
    df=pd.read_csv(filename,header=None,error_bad_lines=False,encoding='utf8')
    df= df.drop(0)
    df.columns=['Date', 'Chat']
    Message= df["Chat"].str.split("-", n = 1, expand = True) 
    df['Date']=df['Date'].str.replace(",","") 
    df['Time']=Message[0]
    df['Text']=Message[1]
    Message1= df["Text"].str.split(":", n = 1, expand = True) 
    df['Text']=Message1[1]
    df['Name']=Message1[0]
    df=df.drop(columns=['Chat'])
    df['Text']=df['Text'].str.lower()
    df['Text'] = df['Text'].str.replace('<media omessi>','**Foto o Video**')
    df['Text'] = df['Text'].str.replace('hai eliminato questo messaggio','Messaggio Eliminato')
    df['Text'] = df['Text'].str.replace('Questo messaggio è stato eliminato','Messaggio Eliminato')    
    df.to_csv("chat.csv",index=False)

def _create_user1_paragraph(d, text, alignment):
    # Add a paragraph
    p = d.add_paragraph()
    p.alignment = 2
    p.paragraph_format.line_spacing = 1.5

    #Paragraph format
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(15)
    paragraph_format.space_after = Pt(3)

    # Add text to paragraph reference
    txt = text
    run = p.add_run(txt)

    # Get the XML tag
    tag = run._r

    # Create XML element
    shd = OxmlElement('w:shd')

    # Add attributes to the element
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '00CC33')

    # Set the font size - this is important! Without this step the
    # tag.rPr value below will be None.
    run.font.size = Pt(10)
    run.font.name = 'Ubuntu'

    tag.rPr.append(shd)

def _create_user2_paragraph(d, text, alignment):
    # Add a paragraph
    p = d.add_paragraph()
    p.alignment = 0
    p.paragraph_format.line_spacing = 1.5

    #Paragraph format
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = Pt(15)
    paragraph_format.space_after = Pt(3)

    # Add text to paragraph reference
    txt = text
    run = p.add_run(txt)

    # Get the XML tag
    tag = run._r

    # Create XML element
    shd = OxmlElement('w:shd')

    # Add attributes to the element
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFFFF')

    # Set the font size - this is important! Without this step the
    # tag.rPr value below will be None.
    run.font.size = Pt(10)
    run.font.name = 'Ubuntu'

    tag.rPr.append(shd)

def _create_document():
    d = Document()
    document = Document()
    section = d.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(50)
    section.right_margin = Mm(50)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

    # Now Add below children to root xml tree
    # create xml element using OxmlElement
    shd = OxmlElement('w:background')
    # Add attributes to the xml element
    shd.set(qn('w:color'), 'F5F4EF') #black color
    shd.set(qn('w:themeColor'), 'text1')
    shd.set(qn('w:themeTint'), 'F2')
    # Add background element at the start of Document.xml using below
    d.element.insert(0,shd)
    # Add displayBackgroundShape element to setting.xml
    shd1 = OxmlElement('w:displayBackgroundShape')
    d.settings.element.insert(0,shd1)

    

    file = csv.reader(open('chat.csv', 'r'))

    user1 = ''
    user2 = ''
    i = 0

    for line in file:
        if i <= 3:
            i = i+1 
            continue
        if user1 == '':
            user1 = line[3]
        if line[3] != user1:
            user2 = line[3]
            break
        
        

    print(user1)
    print(user2)
    
    #composing the document
    d.add_heading('Chat con'+user2, 0)

    i = 0
    for line in tqdm(file, desc="Parsing messages..."):
        if line[3] == user1:
            _create_user1_paragraph(d, line[2]+' - '+line[1], 2)
        elif line[3] == user2:
            _create_user2_paragraph(d,  line[1]+' - '+line[2], 0)
        i = i+1

        #temporary limit in the loop for testing purposes
        # if i == 1000:
        #     break

    d.save('chat.docx')

if not os.path.isfile('./chat.csv'):
    if os.path.isfile('./chat.txt'):
        print('CSV file not found, converting chat.txt in chat.csv..')
        _convert_to_csv(filename)
        print('\nAll ok now, creating document...\n')
        _create_document()
    else:
        print('Cannot find chat.txt nor chat.csv in folder, please add the chat source')
else:
    print('\nAll ok, creating document...\n')
    _create_document()